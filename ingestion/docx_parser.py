from docx import Document
from pathlib import Path
import logging
import re

logger = logging.getLogger(__name__)

def _clean_docx_text(text: str) -> str:
    """
    Pulisce il testo di un documento .docx rimuovendo artefatti comuni.
    """
    # Rimuove spazi eccessivi e linee vuote multiple
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    # Rimuove spazi multipli
    text = re.sub(r' +', ' ', text)
    
    # Rimuove caratteri di controllo non stampabili
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)
    
    return text.strip()

def _extract_paragraph_text(paragraph) -> str:
    """
    Estrae il testo da un paragrafo preservando la formattazione di base.
    """
    text_parts = []
    
    for run in paragraph.runs:
        text = run.text
        
        # Preserva enfasi di base (grassetto, corsivo) con markdown
        if run.bold and run.italic:
            text = f"**_{text}_**"
        elif run.bold:
            text = f"**{text}**"
        elif run.italic:
            text = f"_{text}_"
        
        text_parts.append(text)
    
    return ''.join(text_parts)

def _extract_table_text(table) -> str:
    """
    Estrae il testo dalle tabelle convertendole in formato markdown.
    """
    table_text = []
    
    for i, row in enumerate(table.rows):
        row_text = []
        for cell in row.cells:
            # Pulisce il testo della cella
            cell_text = cell.text.strip().replace('\n', ' ')
            row_text.append(cell_text)
        
        # Crea la riga markdown
        table_text.append('| ' + ' | '.join(row_text) + ' |')
        
        # Aggiunge la riga separatore dopo l'header
        if i == 0:
            separator = '| ' + ' | '.join(['---'] * len(row_text)) + ' |'
            table_text.append(separator)
    
    return '\n'.join(table_text)

def extract_text_from_docx(file_path: Path) -> str:
    """
    Estrae il testo completo da un file .docx preservando struttura e formattazione di base.

    Args:
        file_path: Il percorso del file .docx.

    Returns:
        Una stringa contenente il testo estratto dal documento Word.
        Restituisce una stringa vuota se il file non può essere elaborato.
    """
    try:
        logger.info(f"Inizio estrazione testo dal file DOCX: {file_path}")
        
        doc = Document(file_path)
        content_parts = []
        
        # Estrae i paragrafi principali
        for paragraph in doc.paragraphs:
            text = _extract_paragraph_text(paragraph)
            
            if text.strip():
                # Gestisce gli stili di titolo
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.replace('Heading ', '')) if 'Heading ' in paragraph.style.name else 1
                    text = '#' * level + ' ' + text
                
                content_parts.append(text)
        
        # Estrae le tabelle
        for table in doc.tables:
            table_text = _extract_table_text(table)
            if table_text.strip():
                content_parts.append('\n' + table_text + '\n')
        
        # Unisce tutto il contenuto
        full_text = '\n'.join(content_parts)
        
        # Pulisce il testo finale
        cleaned_text = _clean_docx_text(full_text)
        
        logger.info(f"Estrazione completata per {file_path}. Estratti {len(cleaned_text)} caratteri.")
        return cleaned_text
        
    except Exception as e:
        logger.error(f"Errore durante l'estrazione del testo dal file DOCX {file_path}: {e}")
        return "" 